from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches,Pt

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'couriernew'
font.size = Pt(13)

source = requests.get("http://akul.me/blog/2016/beautifulsoup-cheatsheet/").text
soup = BeautifulSoup(source,'lxml')
t = []
ar = []
for title in soup.find_all('h3'):
        t.append(title.text)
for article in  soup.find_all('code'):
	ar.append(article.text)

dict_ = zip(t,ar)

for key,value in dict_:
        document.add_heading(key, level=0)
        document.add_paragraph(value, style="Quote") 

document.save("bs4_sheet.docx")
